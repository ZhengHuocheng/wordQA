from docx import Document
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import streamlit as st
import json
from docx.shared import Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
# 定义格式检查规则
RULES = {
    '页面大小': ('Letter', 21, 29.7),  # 页面大小（纸张类型、宽度、高度cm）
    '页边距': {
        '左边距': 3.17,  # 单位：cm
        '右边距': 3.17,
        '上边距': 2.54,
        '下边距': 2.54
    },
    '题目': {
        '字体': '黑体',
        '字号': 16,
        '对齐方式': '居中'
    },
    'Heading 1': {
        '字体': '黑体',
        '字号': 10,
        '对齐方式': '左对齐',
    },
    'Heading 2': {
        '字体': '黑体',
        '字号': 12,
        '对齐方式': '左对齐',
    },
    'Heading 3': {
        '字体': '黑体',
        '字号': 12,
        '对齐方式': '左对齐',
    },
    '正文': {
        '字体': '宋体',
        '字号': 10,    #磅
        '首行缩进': 20,  #磅，缩进磅数/字号≈一个字符大小
        '行距': 12    #磅
    }
}
# 创建一个区域，允许用户将文件拖动到其中
uploaded_files = st.file_uploader("将Word文件拖到此处", type=["docx"],accept_multiple_files=True)
# 显示文件路径
doc = None
if uploaded_files is not None:
    for file in uploaded_files:
        st.write(f"文件路径：{file.name}")
        # 打开Word文档
        doc = Document(file.name)
#         print(file.name[:-5])

if doc:
    # 错误信息列表
    errors = []

    # 提取页面大小
    sections = doc.sections
    if sections:
        first_section = sections[0]
        page_width = round(first_section.page_width.cm, 2)  # 将页面宽度转换为cm，并保留两位小数
        page_height = round(first_section.page_height.cm, 2)  # 将页面高度转换为cm，并保留两位小数
        expected_page_size = RULES['页面大小'][1:]
        if (page_width, page_height) != expected_page_size:
            error = f"页面大小为{page_width} x {page_height} cm，不符合规定的{expected_page_size[0]} x {expected_page_size[1]} cm"
            errors.append(error)

    # 提取页边距
    margins = (
        round(first_section.left_margin.cm, 2),  # 将左边距转换为cm，并保留两位小数
        round(first_section.right_margin.cm, 2),  # 将右边距转换为cm，并保留两位小数
        round(first_section.top_margin.cm, 2),  # 将上边距转换为cm，并保留两位小数
        round(first_section.bottom_margin.cm, 2)  # 将下边距转换为cm，并保留两位小数
    )
    expected_margins = RULES['页边距']
    for i, (margin, expected_margin) in enumerate(zip(margins, expected_margins.values()), 1):
        if margin != expected_margin:
            error = f"第{i}个页边距为{margin} cm，不符合规定的{expected_margin} cm"
            errors.append(error)
    title_index=[]
    # 提取题目信息
    title_paragraph = doc.paragraphs[0]
    expected_title_info = RULES['题目']
    title_run = title_paragraph.runs[0]
    title_font_name = title_run.font.name
    title_font_size = title_run.font.size.pt
    title_index.append(0)
    # 提取对齐方式
    title_alignment = str(title_paragraph.alignment)
    if title_alignment == 'CENTER (1)':
        title_alignment = '居中'
    elif title_alignment == 'RIGHT (2)':
        title_alignment = '右对齐'
    elif title_alignment == 'JUSTIFY (3)':
        title_alignment = '两端对齐'
    elif title_alignment == 'LEFT (0)':
        title_alignment = '左对齐'
    if (title_font_name, title_font_size, title_alignment) != (expected_title_info['字体'], expected_title_info['字号'], expected_title_info['对齐方式']):
        error = f"题目的字体为'{title_font_name}'，字号为{title_font_size}，对齐方式为'{title_alignment}'，不符合规定的'{expected_title_info['字形']}'、{expected_title_info['字号']}、'{expected_title_info['对齐方式']}'"
        errors.append(error)

    # 提取各级标题信息
    # 定义各级标题样式名称
    heading_styles = ['Heading 1', 'Heading 2', 'Heading 3']
    # 遍历文档中的段落
    for i,paragraph in enumerate(doc.paragraphs):
        style_name = paragraph.style.name 
        # 判断段落的样式是否为各级标题样式
        if style_name in heading_styles:
            title_index.append(i)
            # 获取标题级别
            heading_level = heading_styles.index(style_name) + 1      
            # 输出标题级别和文本内容
            #print(f"Level {heading_level} Heading: {paragraph.text}")
            # 提取本次标题的字体、字号、对齐方式和行距
            this_title_paragraph = doc.paragraphs[i]
            this_title_run = this_title_paragraph.runs[0]
            this_title_font_name = this_title_run.font.name
            this_title_font_size = this_title_run.font.size.pt
            # 提取对齐方式
            this_title_alignment = str(this_title_paragraph.paragraph_format.alignment)
            if this_title_alignment == 'LEFT (0)':
                this_title_alignment = '左对齐'
            elif this_title_alignment == 'CENTER (1)':
                this_title_alignment = '居中'
            elif this_title_alignment == 'RIGHT (2)':
                this_title_alignment = '右对齐'
            elif this_title_alignment == 'JUSTIFY (3)':
                this_title_alignment = '两端对齐'
            # 与期望的标题信息进行比较并添加错误信息
            expected_this_title_info = RULES[style_name]
            if (this_title_font_name, this_title_font_size, this_title_alignment) != (
                    expected_this_title_info['字体'], expected_this_title_info['字号'],expected_this_title_info['对齐方式']):
                error = f"{style_name}的字体为{this_title_font_name}，字号为{this_title_font_size}，" \
                        f"对齐方式为{this_title_alignment}，" \
                        f"不符合规定的{expected_this_title_info['字体']}、{expected_this_title_info['字号']}、" \
                        f"{expected_this_title_info['对齐方式']}"
                errors.append(error)
    print("题目和各级标题的段落索引为",title_index)
    # 提取正文信息
    for j,paragraph in enumerate(doc.paragraphs):
      if j in title_index:
        continue
      else:
        # print(f"第{j}段")
        expected_font_name = RULES['正文']['字体']
        expected_font_size = RULES['正文']['字号']
        expected_indent = RULES['正文']['首行缩进']
        expected_spacing = RULES['正文']['行距']
        this_paragraph = doc.paragraphs[j]
        this_run = this_paragraph.runs[0]
        this_font_name = this_run.font.name
        this_font_size = this_run.font.size.pt
        # print(this_font_name)
        # print(this_font_size)
        this_paragraph_indent = this_paragraph.paragraph_format.first_line_indent.pt
        this_paragraph_spacing = this_paragraph.paragraph_format.line_spacing.pt
        # print(this_paragraph_indent)
        # print(this_paragraph_spacing)
        if (this_font_name, this_font_size, this_paragraph_indent, this_paragraph_spacing) != (
                expected_font_name, expected_font_size, expected_indent, expected_spacing):
            error = f"第{j+1}段的字体为{this_font_name}，字号为{this_font_size}，首行缩进为{this_paragraph_indent}，行距为{this_paragraph_spacing}，不符合规定的{expected_font_name}、{expected_font_size}、{expected_indent}、{expected_spacing}"
            errors.append(error)

    # 修正错误信息
    # if errors:
    #     print("以下是检测到的格式错误：")
    #     for error in errors:
    # #         print(error)
    error_content = "\n".join(errors) + "."
    error_content = "文档检查错误如下：\n" + error_content
    print(error_content)

    #------------------------------------------------------------------------------------------
    #布局
    import streamlit as st
    from collections import deque
    import json
    # Define question and answer mapping
    qa_pairs = {
        "你好": "你好，有什么可以帮助你的吗？",
        "你叫什么名字？": "我是聊天机器人。",
        "开启修正模式": "修正模式已开启，Worder请注意。下面我会给出设置接口，请给出你的要求",
        "了解":"根据提示进行操作",
        "完成":"好的，请稍等",
        "再见": "再见，祝你有个美好的一天！"
    }
    qa_pairs["当前问题"] = error_content
    DIALOG_HISTORY_FILE = "dialog_history.json"

    # 检查是否有对话历史记录文件存在
    try:
        with open(DIALOG_HISTORY_FILE, "r") as f:
            dialog_history = json.load(f)
    except FileNotFoundError:
        dialog_history = []
#     dialog_history.append(("User", None))
#     dialog_history.append(("ChatBot", error_content))
    # Get answer
    def get_answer(question):
        # Match user input question with predefined questions
        if question in qa_pairs:
            return qa_pairs[question]
        else:
            return "抱歉，我无法回答你的问题。"
    # Set application title
    st.title("聊天系统")
    # Add user input text box
    user_input = st.sidebar.text_input("请输入你的问题")
    # Handle user input and get answer
    if user_input:
        dialog_history.append(("User", user_input))
        if user_input == "了解":
            answer = "根据提示进行输入"
        elif get_answer(user_input):
            answer = get_answer(user_input)
        dialog_history.append(("ChatBot", answer))
    if "设置" in user_input:
        dialog_history.append(("User", user_input))
        user_settings = []
        user_settings = user_input[2:].split(";")
        print("用户输入：",user_input[2:])
        print("接收输入：",user_settings)
        with open("setting.json", "w") as f:
            json.dump(user_settings, f)
        #用户自定义设置信息,格式"21,29;3.17,3,17,2.57,2.57;黑体,16,居中;黑体,10,左对齐;黑体,12,左对齐;黑体,12,左对齐;宋体,10,20,12"
        answer = "收到"
        dialog_history.append(("ChatBot", answer))
    # 将更新后的对话历史记录保存到文件中
    with open(DIALOG_HISTORY_FILE, "w") as f:
        json.dump(dialog_history, f)
#     print(dialog_history)
    # Display dialog history in the sidebar
    st.sidebar.subheader("当前对话")
    for i, (role, message) in enumerate(dialog_history):
        if(i == len(dialog_history)-1):
            if role == "User":
                st.sidebar.text_area(f"User-{i}", value=message, height=20,disabled=True)
            else:
                st.sidebar.text_area(f"ChatBot-{i}", value=message, height=20,disabled=True)
    # Display current dialog in the main area
    for i, (role, message) in enumerate(dialog_history):
        if role == "User":
            st.text_area(f"User-{100+i}", value=message, height=20,disabled=True)
        else:
            st.text_area(f"ChatBot-{100+i}", value=message, height=20,disabled=True)
    # --------------------------------------------------------------------------------------------

    # 修正模式

    #-------------------------------------------------------------------------------------------------
    if user_input == "完成":
        try:
            with open("setting.json", "r") as f:
                user_settings = json.load(f)
        except FileNotFoundError:
                user_settings = []
        # # 执行修正操作
        settings_split_list = user_settings
        #['21,29', '3,17,3,17,2,2', '黑体,16,居中', '黑体,10,左对齐', '黑体,12,左对齐', '黑体,12左对齐', '宋体,10,20,12']
        pagesize = settings_split_list[0]
        margin = settings_split_list[1]
        titlesetting = settings_split_list[2]
        Heading_1 = settings_split_list[3]
        Heading_2 = settings_split_list[4]
        Heading_3 = settings_split_list[5]
        main_pargraph = settings_split_list[6]
        #按,分裂
        pagesize_split = pagesize.split(",")
        margin_split = margin.split(",")
        titlesetting_split = titlesetting.split(",")
        Heading_1_split = Heading_1.split(",")
        Heading_2_split = Heading_2.split(",")
        Heading_3_split = Heading_3.split(",")
        main_pargraph_split = main_pargraph.split(",")
        # # 提取页面大小规定
        # expected_page_size = RULES['页面大小']
        # 设置页面大小
        sections = doc.sections
        if sections:
          for i,section in enumerate(sections):
            section = sections[i]
            page_width = Cm(float(pagesize_split[0]))  # 将宽度转换为Cm单位
            page_height = Cm(float(pagesize_split[1]))  # 将高度转换为Cm单位
            section.page_width = page_width
            section.page_height = page_height

        # 设置页边距规格
        # expected_margins = RULES['页边距']
        if sections:
          for i,section in enumerate(sections):
            section = sections[i]
            # 左边距
            if round(section.left_margin.cm, 2) != float(margin_split[0]):
                section.left_margin = Cm(float(margin_split[0]))
            # 右边距
            if round(section.right_margin.cm, 2) != float(margin_split[1]):
                section.right_margin = Cm(float(margin_split[1]))
            # 上边距
            if round(section.top_margin.cm, 2) != float(margin_split[2]):
                section.top_margin = Cm(float(margin_split[2]))
            # 下边距
            if round(section.bottom_margin.cm, 2) != float(margin_split[3]):
                section.bottom_margin = Cm(float(margin_split[3]))
        # # 设置题目样式
        title_paragraph = doc.paragraphs[0]
        # expected_title_info = RULES['题目']
        # # 设置题目字体
        title_run = title_paragraph.runs[0]
        title_run.font.name = titlesetting_split[0]
        title_run.font.size = Pt(float(titlesetting_split[1]))
        # 设置对齐方式
        title_alignment = titlesetting_split[2]
        if title_alignment == '居中':
            title_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        elif title_alignment == '右对齐':
            title_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        elif title_alignment == '两端对齐':
            title_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        elif title_alignment == '左对齐':
            title_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        #设置各级标题
        # 遍历文档中的段落
        for i, paragraph in enumerate(doc.paragraphs):
            style_name = paragraph.style.name
            # 判断段落的样式是否为各级标题样式
            if style_name in heading_styles:
                # 提取本次标题的字体、字号、对齐方式和行距
                if style_name == 'Heading 1':
                    HeadingSplit = Heading_1_split
                if style_name == 'Heading 2':
                    HeadingSplit = Heading_2_split
                if style_name == 'Heading 3':
                    HeadingSplit = Heading_3_split
                this_title_paragraph = doc.paragraphs[i]
#                 this_title_run = this_title_paragraph.runs[0]
                for run in this_title_paragraph.runs:
                    run.font.name = expected_font_name
                    run.font.size = expected_font_size
                this_title_run.font.name = HeadingSplit[0]
                this_title_run.font.size = Pt(int(HeadingSplit[1]))
                # 提取对齐方式
                if HeadingSplit[2] == '左对齐':
                    this_title_paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                elif HeadingSplit[2] == '居中':
                    this_title_paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                elif HeadingSplit[2] == '右对齐':
                    this_title_paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                elif HeadingSplit[2] == '两端对齐':
                    this_title_paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        #设置正文样式
        for j, this_paragraph in enumerate(doc.paragraphs):
            if j in title_index:
                continue
            else:
                expected_font_name = main_pargraph_split[0]
                expected_font_size = Pt(int(main_pargraph_split[1]))
                expected_indent = Pt(int(main_pargraph_split[2]))
                expected_spacing = Pt(int(main_pargraph_split[3]))
                this_run = this_paragraph.runs[0]
                for run in this_paragraph.runs:
                    run.font.name = expected_font_name
                    run.font.size = expected_font_size
                this_paragraph.paragraph_format.first_line_indent = expected_indent
                this_paragraph.paragraph_format.line_spacing = expected_spacing
#         保存修正后的文档
        doc.save(file.name[:-5]+"_copy.docx")